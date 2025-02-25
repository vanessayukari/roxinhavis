import matplotlib
from flask import Flask, render_template, request, send_file
import pandas as pd
import os
import matplotlib.pyplot as plt
from docx import Document
from io import BytesIO
import zipfile
import requests

app = Flask(__name__)

# Logo para barra lateral
def baixar_imagens(df, coluna_logo, pasta='static/img/streamers'):
    if not os.path.exists(pasta):
        os.makedirs(pasta)

    caminhos_imagens = []

    for _, row in df.iterrows():
        url_imagem = row[coluna_logo]

        if url_imagem and isinstance(url_imagem, str):
            try:
                resposta = requests.get(url_imagem, stream=True)
                if resposta.status_code == 200:
                    nome_imagem = os.path.join(pasta, f"{row['Channel']}_logo.png")
                    with open(nome_imagem, 'wb') as f:
                        for chunk in resposta.iter_content(1024):
                            f.write(chunk)
                    caminhos_imagens.append(nome_imagem)
                else:
                    print(f"Erro ao baixar a imagem para {row['Channel']}: {resposta.status_code}")
            except Exception as e:
                print(f"Erro ao tentar baixar a imagem de {row['Channel']}: {e}")

    return caminhos_imagens

# Gerar análise
def gerar_analise(file1, file2, mes1, ano1, mes2, ano2, jogo):
    df_mes1 = pd.read_csv(file1)
    df_mes2 = pd.read_csv(file2)

    logos_mes1 = baixar_imagens(df_mes1, 'Logo')
    logos_mes2 = baixar_imagens(df_mes2, 'Logo')

    # Filtrar
    df_comparacao = pd.merge(df_mes1, df_mes2, on="Channel", suffixes=('_mes1', '_mes2'))

    # Crescimento de seguidores
    df_comparacao['crescimento_seguidores'] = df_comparacao['Followers_mes2'] - df_comparacao['Followers_mes1']
    df_comparacao['porcentagem_crescimento_seguidores'] = (
            (df_comparacao['crescimento_seguidores'] / df_comparacao['Followers_mes1']) * 100
    ).fillna(0)

    # Tempo assistido (WT)
    df_comparacao['tempo_assistido'] = df_comparacao['Watch time (mins)_mes2'] - df_comparacao['Watch time (mins)_mes1']
    df_comparacao['tempo_assistido'] = df_comparacao['tempo_assistido'] / 60

    df_comparacao['porcentagem_tempo_assistido'] = 0
    df_comparacao.loc[df_comparacao['Watch time (mins)_mes1'] > 0, 'porcentagem_tempo_assistido'] = (
            (df_comparacao['tempo_assistido'] / (df_comparacao['Watch time (mins)_mes1'] / 60)) * 100
    )

    df_comparacao['porcentagem_tempo_assistido'] = df_comparacao['porcentagem_tempo_assistido'].fillna(0)

    # Tempo streamado (ST)
    df_comparacao['tempo_streamado'] = df_comparacao['Stream time (mins)_mes2'] - df_comparacao[ 'Stream time (mins)_mes1']
    df_comparacao['tempo_streamado'] = df_comparacao['tempo_streamado'] / 60

    df_comparacao['porcentagem_tempo_streamado'] = 0
    df_comparacao.loc[df_comparacao['Stream time (mins)_mes1'] > 0, 'porcentagem_tempo_streamado'] = (
            (df_comparacao['tempo_streamado'] / (df_comparacao['Stream time (mins)_mes1'] / 60)) * 100
    )

    df_comparacao['porcentagem_tempo_streamado'] = df_comparacao['porcentagem_tempo_streamado'].fillna(0)

    # Análises
    insights = []
    insights.append(f"Análise de criadores de: {jogo}")
    insights.append(f"Meses comparados: {mes1}/{ano1} e {mes2}/{ano2}")

    crescimento = []
    for _, row in df_comparacao[df_comparacao['crescimento_seguidores'] > 0].iterrows():
        crescimento.append(
            f"<strong>{row['Channel']}</strong> teve um crescimento de {row['crescimento_seguidores']:,.2f} seguidores ({row['porcentagem_crescimento_seguidores']:.2f}%).")

    queda = []
    for _, row in df_comparacao[df_comparacao['crescimento_seguidores'] < 0].iterrows():
        queda.append(
            f"<strong>{row['Channel']}</strong> teve uma queda de {abs(row['crescimento_seguidores']):,.2f} seguidores ({row['porcentagem_crescimento_seguidores']:.2f}%).")

    crescimento_tempo_assistido = []
    for _, row in df_comparacao[df_comparacao['tempo_assistido'] > 0].iterrows():
        crescimento_tempo_assistido.append(
            f"<strong>{row['Channel']}</strong> teve um crescimento de {row['tempo_assistido']:,.2f} horas assistidas ({row['porcentagem_tempo_assistido']:.2f}%).")

    crescimento_tempo_streamado = []
    for _, row in df_comparacao[df_comparacao['tempo_streamado'] > 0].iterrows():
        crescimento_tempo_streamado.append(
            f"<strong>{row['Channel']}</strong> teve um crescimento de {row['tempo_streamado']:,.2f} horas streamadas ({row['porcentagem_tempo_streamado']:.2f}%).")

    sumidos = df_mes1[~df_mes1['Channel'].isin(df_mes2['Channel'])]
    sumidos_lista = []
    if not sumidos.empty:
        insights.append('Streamers que não aparecem no Top 50:')
        for _, row in sumidos.iterrows():
            sumidos_lista.append(f"<strong>{row['Channel']}</strong> não apareceu em {mes2}/{ano2}")

    # Top 10
    top_10_mes1 = df_mes1.nlargest(10, 'Watch time (mins)')
    top_10_mes1_lista = [
        f"<strong>{row['Channel']}</strong> - {row['Watch time (mins)'] / 60:,.2f} horas" for _, row in top_10_mes1.iterrows()
    ]

    top_10_mes2 = df_mes2.nlargest(10, 'Watch time (mins)')
    top_10_mes2_lista = [
        f"<strong>{row['Channel']}</strong> - {row['Watch time (mins)'] / 60:,.2f} horas" for _, row in top_10_mes2.iterrows()
    ]

    # Função para gerar gráficos
    cores = ['#5245FF', '#CC45FF', '#4573FF', '#FF45EC', '#C29AFF', '#FF457D', '#FF5545', '#FFA29A', '#FF9645', '#8E45FF']
    def grafico_top_10(top_10_df, mes, ano):
        canais = top_10_df['Channel']
        tempo_assistido = top_10_df['Watch time (mins)'] / 60

        pasta_imagens = os.path.join('static', 'img')
        if not os.path.exists(pasta_imagens):
            os.makedirs(pasta_imagens)

        matplotlib.rcParams['font.family'] = 'Segoe UI'

        plt.figure(figsize=(8, 8))
        plt.pie(tempo_assistido, labels=canais, autopct='%1.1f%%', startangle=140, colors=cores)

        filepath = os.path.join(pasta_imagens, f"top_10_{mes}_{ano}.png")
        plt.savefig(filepath, bbox_inches='tight')
        plt.close()

        return f"img/top_10_{mes}_{ano}.png"

    grafico_mes1_path = grafico_top_10(top_10_mes1, mes1, ano1)
    grafico_mes2_path = grafico_top_10(top_10_mes2, mes2, ano2)

    return {
        'crescimento': crescimento,
        'queda': queda,
        'crescimento_tempo_assistido': crescimento_tempo_assistido,
        'crescimento_tempo_streamado': crescimento_tempo_streamado,
        'sumidos': sumidos_lista,
        'top_10_mes1': top_10_mes1_lista,
        'top_10_mes2': top_10_mes2_lista,
        'grafico_mes1': grafico_mes1_path,
        'grafico_mes2': grafico_mes2_path
    }

# Gerar docx
def gerar_docx(insights, jogo, mes1, ano1, mes2, ano2):
    doc = Document()
    doc.add_heading(f"Análise de Criadores de: {jogo}", 0)
    doc.add_paragraph(f"Meses comparados: {mes1}/{ano1} e {mes2}/{ano2}")

    doc.add_heading('Crescimento de Seguidores:', level=1)
    for item in insights['crescimento']:
        doc.add_paragraph(item)

    doc.add_heading('Queda de Seguidores:', level=1)
    for item in insights['queda']:
        doc.add_paragraph(item)

    doc.add_heading('Crescimento de Tempo Assistido:', level=1)
    for item in insights['crescimento_tempo_assistido']:
        doc.add_paragraph(item)

    doc.add_heading('Crescimento de Tempo Streamado:', level=1)
    for item in insights['crescimento_tempo_streamado']:
        doc.add_paragraph(item)

    if insights['sumidos']:
        doc.add_heading('Streamers que Não Aparecem no Top 50:', level=1)
        for item in insights['sumidos']:
            doc.add_paragraph(item)

    doc.add_heading(f'Top 10 - {mes1}/{ano1}:', level=1)
    for item in insights['top_10_mes1']:
        doc.add_paragraph(item)

    doc.add_heading(f'Top 10 - {mes2}/{ano2}:', level=1)
    for item in insights['top_10_mes2']:
        doc.add_paragraph(item)

    file_path = f"resultados_{jogo}_{mes1}_{ano1}_vs_{mes2}_{ano2}.docx"
    doc.save(file_path)
    return file_path

# Rota
@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        file1 = request.files['csv_mes1']
        file2 = request.files['csv_mes2']
        mes1 = request.form['mes1']
        ano1 = request.form['ano1']
        mes2 = request.form['mes2']
        ano2 = request.form['ano2']
        jogo = request.form['jogo']

        file1_path = os.path.join('uploads', file1.filename)
        file2_path = os.path.join('uploads', file2.filename)

        file1.save(file1_path)
        file2.save(file2_path)

        insights = gerar_analise(file1_path, file2_path, mes1, ano1, mes2, ano2, jogo)

        doc_path = gerar_docx(insights, jogo, mes1, ano1, mes2, ano2)

        return render_template('resultados.html', insights=insights, jogo=jogo, mes1=mes1, ano1=ano1, mes2=mes2, ano2=ano2, doc_path=doc_path)

    return render_template('index.html')

@app.route('/download/<filename>')
def download(filename):
    mes1 = request.args.get('mes1')
    ano1 = request.args.get('ano1')
    mes2 = request.args.get('mes2')
    ano2 = request.args.get('ano2')

    grafico_mes1 = os.path.join('static', 'img', f"top_10_{mes1}_{ano1}.png")
    grafico_mes2 = os.path.join('static', 'img', f"top_10_{mes2}_{ano2}.png")
    doc_path = filename

    if not os.path.exists(grafico_mes1) or not os.path.exists(grafico_mes2):
        return "Erro: gráficos não encontrados."

    memory_file = BytesIO()

    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
        zipf.write(doc_path, os.path.basename(doc_path))
        zipf.write(grafico_mes1, os.path.basename(grafico_mes1))
        zipf.write(grafico_mes2, os.path.basename(grafico_mes2))

    memory_file.seek(0)
    return send_file(memory_file, as_attachment=True, download_name="resultados.zip", mimetype="application/zip")

if __name__ == '__main__':
    app.run(debug=True)
